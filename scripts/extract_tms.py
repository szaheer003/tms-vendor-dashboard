#!/usr/bin/env python3
"""
FIS TMS vendor workbook + brief extraction. Re-run to regenerate /public/data/*.json
"""
from __future__ import annotations

import json
import math
import os
import re
import shutil
import sys
from collections import defaultdict
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

import openpyxl
from openpyxl.utils import column_index_from_string, get_column_letter
from docx import Document

_SCRIPT_DIR = Path(__file__).resolve().parent
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))
from folder2_pdf_resolver import proposal_paths_for_extraction

BASE = Path(os.environ.get("TMS_BASE", Path(__file__).resolve().parents[1]))
F1 = BASE / "Folder 1"
F2 = BASE / "Folder 2"
F3 = BASE / "Folder 3"
F4 = BASE / "Folder 4"
OUT = BASE / "src" / "data"
EXTRACTION_TIMESTAMP = datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def parse_a1(ref: str) -> Optional[tuple[int, int]]:
    m = re.match(r"^([A-Z]+)(\d+)$", ref.strip().upper().replace("$", ""))
    if not m:
        return None
    return int(m.group(2)), column_index_from_string(m.group(1))


def grid_snapshot(
    ws,
    min_row: int,
    max_row: int,
    min_col: int,
    max_col: int,
    highlights: set[tuple[int, int]],
) -> dict[str, Any]:
    """2D array of cells for mini spreadsheet preview in UI."""
    rows_out: list[list[dict[str, Any]]] = []
    for r in range(min_row, max_row + 1):
        row_cells: list[dict[str, Any]] = []
        for c in range(min_col, max_col + 1):
            v = ws.cell(r, c).value
            if isinstance(v, float) and math.isnan(v):
                disp = ""
            elif v is None:
                disp = ""
            elif isinstance(v, (int, float)) and abs(v) > 1e6:
                disp = f"{v:.10g}"[:14]
            else:
                s = str(v).strip().replace("\n", " ")
                disp = s[:48] + ("…" if len(s) > 48 else "")
            row_cells.append(
                {
                    "r": r,
                    "c": c,
                    "col": get_column_letter(c),
                    "value": disp,
                    "highlight": (r, c) in highlights,
                }
            )
        rows_out.append(row_cells)
    return {
        "minRow": min_row,
        "maxRow": max_row,
        "minCol": min_col,
        "maxCol": max_col,
        "rows": rows_out,
    }


def snapshot_around_row(ws, center_row: int, min_col: int = 2, max_col: int = 12) -> dict[str, Any]:
    """Context ±3 rows around center; highlight quarterly band C:V on center row."""
    mr = max(1, center_row - 3)
    xr = min((ws.max_row or center_row) + 1, center_row + 3)
    mc = min_col
    xc = min(max_col, ws.max_column or max_col)
    hi = {(center_row, c) for c in range(3, 23) if c <= xc}
    return grid_snapshot(ws, mr, xr, mc, xc, hi)


def cell_snapshot(ws, ref: str, pad: int = 3) -> Optional[dict[str, Any]]:
    parsed = parse_a1(ref)
    if not parsed:
        return None
    r, c = parsed
    mr, xr = max(1, r - pad), min((ws.max_row or r) + pad, r + pad)
    mc = max(1, c - pad)
    xc = min((ws.max_column or c) + pad, c + pad)
    return grid_snapshot(ws, mr, xr, mc, xc, {(r, c)})


def workbook_meta(
    *,
    source_file: str,
    tab: str,
    location: str,
    value_label: str,
    calculation: str,
    snapshot: Optional[dict[str, Any]],
    verified: bool = True,
) -> dict[str, Any]:
    return {
        "kind": "workbook",
        "sourceFile": source_file,
        "tab": tab,
        "location": location,
        "valueLabel": value_label,
        "calculation": calculation,
        "extractionTimestamp": EXTRACTION_TIMESTAMP,
        "verified": verified,
        "snapshot": snapshot,
    }


def proposal_meta(
    *,
    source_file: str,
    page: str,
    section: str,
    excerpt: str,
    highlight_phrase: str,
) -> dict[str, Any]:
    return {
        "kind": "proposal",
        "sourceFile": source_file,
        "page": page,
        "section": section,
        "excerpt": excerpt,
        "highlightPhrase": highlight_phrase,
        "vendorClaimNote": "Vendor claim — not independently verified.",
        "extractionTimestamp": EXTRACTION_TIMESTAMP,
    }


def ubiquity_price_provenance(
    source_file: str, ws, tcv_m: float, years: list[Optional[float]]
) -> dict[str, Any]:
    hi = set()
    for r in (110, 111, 112):
        for c in range(3, 23):
            hi.add((r, c))
    xr = min(115, ws.max_row or 115)
    snap = grid_snapshot(ws, 105, xr, 2, 13, hi)
    prov: dict[str, Any] = {
        "tcvM": workbook_meta(
            source_file=source_file,
            tab="6.0  Pricing",
            location="Rows 110+111 (quarters); row 112 annual check",
            value_label=f"${tcv_m:.2f}M",
            calculation="Summed quarterly NA+EMEA lines; USD → $M",
            snapshot=snap,
        ),
    }
    for i, yv in enumerate(years[:5]):
        if yv is not None:
            prov[f"year{i+1}"] = workbook_meta(
                source_file=source_file,
                tab="6.0  Pricing",
                location="Derived from rows 110+111 quarterly bands",
                value_label=f"${yv:.2f}M",
                calculation="Four quarters per year from combined rows",
                snapshot=snap,
            )
    return prov


def build_pricing_provenance(
    source_file: str,
    ws,
    total_row: int,
    tcv_m: float,
    years: list[Optional[float]],
    unit_note: str,
) -> dict[str, Any]:
    snap = snapshot_around_row(ws, total_row)
    prov: dict[str, Any] = {
        "tcvM": workbook_meta(
            source_file=source_file,
            tab="6.0  Pricing",
            location=f"Row {total_row}, columns C:V (20 quarters)",
            value_label=f"${tcv_m:.2f}M",
            calculation=f"Sum of quarterly values on Total Annual Cost row; {unit_note}",
            snapshot=snap,
        ),
    }
    for i, yv in enumerate(years[:5]):
        if yv is not None:
            prov[f"year{i+1}"] = workbook_meta(
                source_file=source_file,
                tab="6.0  Pricing",
                location=f"Row {total_row}, Y{i+1} = sum of quarters {i*4+1}–{(i+1)*4}",
                value_label=f"${yv:.2f}M",
                calculation="Sum of four quarterly cells on same row",
                snapshot=snap,
            )
    return prov


def efficiency_snapshot(ws) -> dict[str, Any]:
    last_col = min(14, ws.max_column or 14)
    last_row = min(13, ws.max_row or 13)
    return grid_snapshot(ws, 5, last_row, 2, last_col, set())


def enrich_drill_snippets(wb: openpyxl.Workbook, blocks: list[dict[str, Any]], source_file: str) -> None:
    for block in blocks:
        if block.get("missing"):
            continue
        tab = block.get("tab", "")
        if tab.startswith("Q: "):
            sname = tab[3:].strip()
            if sname not in wb.sheetnames:
                continue
            ws = wb[sname]
        else:
            if tab not in wb.sheetnames:
                continue
            ws = wb[tab]
        for sn in block.get("snippets") or []:
            ref = sn.get("ref", "")
            snap = cell_snapshot(ws, ref)
            sn["snapshot"] = snap
            sn["sourcePreview"] = workbook_meta(
                source_file=source_file,
                tab=tab,
                location=f"Cell {ref}",
                value_label="Vendor response cell",
                calculation="Q/A extraction: template question vs vendor answer column",
                snapshot=snap,
            )


def _needle_norm(s: str, max_len: int = 400) -> str:
    return " ".join((s or "").split()).lower()[:max_len]


def find_text_page_in_pdf(pdf_path: Path, text: str) -> Optional[int]:
    """Return 1-based page index if a substring of normalized answer text appears in the PDF."""
    try:
        import pdfplumber
    except ImportError:
        return None
    if not pdf_path.is_file():
        return None
    raw = _needle_norm(text, 500)
    if len(raw) < 20:
        return None
    for n in (140, 100, 70, 45):
        frag = raw[:n]
        if len(frag) < 20:
            break
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    pt = _needle_norm(page.extract_text() or "", 200000)
                    if frag in pt:
                        return i + 1
        except Exception:
            return None
    return None


def find_text_in_docx(docx_path: Path, text: str) -> bool:
    if not docx_path.is_file():
        return False
    raw = _needle_norm(text, 400)
    if len(raw) < 20:
        return False
    frag = raw[:min(150, len(raw))]
    try:
        doc = Document(str(docx_path))
        for p in doc.paragraphs:
            if frag in _needle_norm(p.text or "", 4000):
                return True
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if frag in _needle_norm(cell.text or "", 4000):
                        return True
    except Exception:
        return False
    return False


def proposal_link_meta(
    *,
    source_file: str,
    page: str,
    section: str,
    excerpt: str,
    highlight: str = "",
    submission_doc: str = "proposal",
    proposal_part: Optional[int] = None,
    supplemental_index: Optional[int] = None,
    search_query: Optional[str] = None,
) -> dict[str, Any]:
    out: dict[str, Any] = {
        "kind": "proposal",
        "sourceFile": source_file,
        "page": page,
        "section": section,
        "excerpt": excerpt[:2000],
        "highlightPhrase": (highlight or "")[:500],
        "vendorClaimNote": "Matched text in vendor submission — confirm in viewer.",
        "extractionTimestamp": EXTRACTION_TIMESTAMP,
        "submissionDoc": submission_doc,
    }
    if proposal_part is not None:
        out["proposalPart"] = proposal_part
    if supplemental_index is not None:
        out["supplementalIndex"] = supplemental_index
    if search_query:
        out["searchQuery"] = search_query[:400]
    return out


_VENDOR_SOW_PATHS: dict[str, Path] = {
    "cognizant": F3 / "Cognizant_Appendix C - Draft FIS SOW_Cognizant Redline.Docx",
    "genpact": F3 / "Genpact_Appendix C - Draft FIS SOW_G updated.docx",
    "exl": F3 / "EXL_Draft FIS SOW - EXL Comment.docx",
    "sutherland": F3 / "Sutherland_Response to Appendix C - FIS SOW.docx",
    "ubiquity": F3 / "Ubiquity_Appendix C - Draft FIS SOW RF redliines CB.docx",
    "ibm": F3 / "IBM_Appendix C - Draft FIS SOW_IBM March 26.docx",
}


def vendor_link_sources_for(vendor_id: str) -> dict[str, Any]:
    """Folder 2 proposal PDFs (resolved dynamically) + Folder 3 SOW paths for drill-down text matching."""
    pdfs = proposal_paths_for_extraction(BASE, vendor_id)
    return {
        "proposal_pdfs": pdfs,
        "sow": _VENDOR_SOW_PATHS.get(vendor_id),
    }


def enrich_drill_linked_documents(vendor_id: str, blocks: list[dict[str, Any]]) -> None:
    """Attach linkedDocumentPreview (proposal PDF or SOW) when answer text matches submission files (requires Folder 2/3)."""
    if os.environ.get("TMS_SKIP_LINKED_PDF", "").strip().lower() in ("1", "true", "yes"):
        return
    cfg = vendor_link_sources_for(vendor_id)
    proposal_list: list[tuple[Path, str, Optional[int]]] = list(cfg.get("proposal_pdfs") or [])
    sow_path: Optional[Path] = cfg.get("sow")
    if not proposal_list and not sow_path:
        return
    for block in blocks:
        if block.get("missing"):
            continue
        tab = str(block.get("tab", ""))
        for sn in block.get("snippets") or []:
            if sn.get("linkedDocumentPreview"):
                continue
            text = (sn.get("text") or "").strip()
            if len(text) < 25:
                continue
            excerpt = text[:600] + ("…" if len(text) > 600 else "")
            hw = text[:80].strip()
            matched = False
            for pdf_path, display_name, part in proposal_list:
                pg = find_text_page_in_pdf(pdf_path, text)
                if pg is not None:
                    sn["linkedDocumentPreview"] = proposal_link_meta(
                        source_file=display_name,
                        page=str(pg),
                        section=f"{tab} → proposal PDF (text match)",
                        excerpt=excerpt,
                        highlight=hw,
                        submission_doc="proposal",
                        proposal_part=part,
                        search_query=text[:350],
                    )
                    matched = True
                    break
            if matched:
                continue
            if sow_path and isinstance(sow_path, Path) and find_text_in_docx(sow_path, text):
                sn["linkedDocumentPreview"] = proposal_link_meta(
                    source_file=sow_path.name,
                    page="—",
                    section=f"{tab} → SOW redline (text match)",
                    excerpt=excerpt,
                    highlight=hw,
                    submission_doc="sow",
                    search_query=text[:350],
                )


def rate_cell_display(v: Any) -> str:
    """String for rate card cells: 2 decimal places for numeric rates."""
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return ""
    if isinstance(v, (int, float)):
        return f"{float(v):.2f}"
    return cell_str(v)


def rate_rows_with_snapshots(ws, source_file: str, r_start: int = 10, r_end: int = 35) -> list[dict[str, Any]]:
    """Section 6.1: include Tier sub-rows (col C) and normalize numeric formatting."""
    rows: list[dict[str, Any]] = []
    pending_tier = ""
    _xc = min(12, ws.max_column or 12)
    for r in range(r_start, r_end + 1):
        b = cell_str(ws.cell(r, 2).value)
        c = cell_str(ws.cell(r, 3).value)
        if b.lower().startswith("section"):
            continue
        if b in ("Service", "Annual COLA Assumption", "Price/Unit ($)"):
            continue
        if b and any(x in b.lower() for x in ("tier 1", "tier 2", "back office")):
            pending_tier = b
        label = b
        if not label and c:
            label = f"{pending_tier} — {c}" if pending_tier else c
        elif b and c and b != c and c.lower() not in ("price/unit ($)",):
            if pending_tier and b == pending_tier:
                label = f"{pending_tier} — {c}"
            elif "tier" not in b.lower():
                label = f"{b} — {c}"
            else:
                label = f"{b} — {c}"
        if not label:
            continue
        has_num = False
        for col in range(4, _xc + 1):
            v = ws.cell(r, col).value
            if isinstance(v, (int, float)) and not (isinstance(v, float) and math.isnan(v)) and abs(float(v)) > 0:
                has_num = True
                break
        if not has_num:
            if not any(cell_str(ws.cell(r, col).value) for col in range(4, _xc + 1)):
                continue
        snap = grid_snapshot(ws, max(1, r - 1), r + 1, 2, _xc, {(r, c) for c in range(4, _xc + 1)})
        rows.append(
            {
                "label": label,
                "row": r,
                "onshore": rate_cell_display(ws.cell(r, 4).value) or cell_str(ws.cell(r, 4).value),
                "nearshore": rate_cell_display(ws.cell(r, 5).value) or cell_str(ws.cell(r, 5).value),
                "offshore": rate_cell_display(ws.cell(r, 6).value) or cell_str(ws.cell(r, 6).value),
                "emeaOnshore": rate_cell_display(ws.cell(r, 7).value) or cell_str(ws.cell(r, 7).value),
                "sourcePreview": workbook_meta(
                    source_file=source_file,
                    tab="6.0  Pricing",
                    location=f"Section 6.1 rate card, row {r}",
                    value_label=label,
                    calculation="Hourly rate as submitted in workbook",
                    snapshot=snap,
                ),
            }
        )
    return rows


def _find_row_col3_contains(ws, r0: int, r1: int, needle: str) -> Optional[int]:
    n = needle.lower()
    for r in range(r0, r1 + 1):
        if n in cell_str(ws.cell(r, 3).value).lower():
            return r
    return None


def _patch_cognizant_cs_live(rows: list[dict[str, Any]], ws, source_file: str) -> None:
    r = _find_row_col3_contains(ws, 11, 22, "multilingual") or 13
    snap = grid_snapshot(ws, max(1, r - 1), r + 1, 2, min(12, ws.max_column or 12), {(r, c) for c in range(4, 11)})
    canon = {
        "label": "Tier 1 Customer Service — Multilingual (CS Live Agent reference)",
        "row": r,
        "onshore": rate_cell_display(ws.cell(r, 4).value),
        "nearshore": rate_cell_display(ws.cell(r, 6).value),
        "offshore": rate_cell_display(ws.cell(r, 7).value),
        "emeaOnshore": rate_cell_display(ws.cell(r, 9).value),
        "sourcePreview": workbook_meta(
            source_file=source_file,
            tab="6.0  Pricing",
            location=f"Section 6.1 — CS live reference row {r} (US col D, DR col F, PH col G, UK col I)",
            value_label="CS Live Agent (canonical columns)",
            calculation="Mapped from workbook geography columns on Multilingual row",
            snapshot=snap,
        ),
    }
    for i, row in enumerate(rows):
        if "multilingual" in row.get("label", "").lower():
            rows[i] = canon
            return
    rows.insert(0, canon)


def _patch_genpact_cs_live(rows: list[dict[str, Any]], ws, source_file: str) -> None:
    r = _find_row_col3_contains(ws, 11, 22, "live agent") or 13
    snap = grid_snapshot(ws, max(1, r - 1), r + 1, 2, min(12, ws.max_column or 12), {(r, c) for c in range(4, 11)})
    near = ws.cell(r, 5).value
    near_s = "" if near in (0, None) or (isinstance(near, float) and math.isnan(near)) else rate_cell_display(near)
    canon = {
        "label": "Tier 1 Customer Service — Live Agent (canonical)",
        "row": r,
        "onshore": rate_cell_display(ws.cell(r, 4).value),
        "nearshore": near_s,
        "offshore": rate_cell_display(ws.cell(r, 6).value),
        "emeaOnshore": rate_cell_display(ws.cell(r, 7).value),
        "sourcePreview": workbook_meta(
            source_file=source_file,
            tab="6.0  Pricing",
            location=f"Section 6.1 row {r} (Genpact US/PH/UK columns)",
            value_label="CS Live Agent",
            calculation="Live Agent sub-row per workbook layout",
            snapshot=snap,
        ),
    }
    for i, row in enumerate(rows):
        if "live agent" in row.get("label", "").lower():
            rows[i] = canon
            return
    rows.insert(0, canon)


def _patch_exl_cs_live(rows: list[dict[str, Any]], ws, source_file: str) -> None:
    r = _find_row_col3_contains(ws, 11, 22, "live agent") or 13
    snap = grid_snapshot(ws, max(1, r - 1), r + 1, 2, min(12, ws.max_column or 12), {(r, c) for c in range(4, 11)})
    canon = {
        "label": "Tier 1 Customer Service — Live Agent (canonical)",
        "row": r,
        "onshore": "",
        "nearshore": "",
        "offshore": rate_cell_display(ws.cell(r, 6).value),
        "emeaOnshore": "",
        "sourcePreview": workbook_meta(
            source_file=source_file,
            tab="6.0  Pricing",
            location=f"Section 6.1 row {r} (EXL offshore India/PH column)",
            value_label="CS Live Agent (offshore)",
            calculation="Live Agent row — primary rate in col F",
            snapshot=snap,
        ),
    }
    for i, row in enumerate(rows):
        if "live agent" in row.get("label", "").lower():
            rows[i] = canon
            return
    rows.insert(0, canon)


def _patch_sutherland_cs_live(rows: list[dict[str, Any]], ws, source_file: str) -> None:
    r = _find_row_col3_contains(ws, 11, 22, "live agent") or 13
    snap = grid_snapshot(ws, max(1, r - 1), r + 1, 2, min(12, ws.max_column or 12), {(r, c) for c in range(4, 11)})
    canon = {
        "label": "Tier 1 Customer Service — Live Agent (canonical)",
        "row": r,
        "onshore": rate_cell_display(ws.cell(r, 4).value),
        "nearshore": rate_cell_display(ws.cell(r, 7).value),
        "offshore": rate_cell_display(ws.cell(r, 6).value),
        "emeaOnshore": rate_cell_display(ws.cell(r, 9).value),
        "sourcePreview": workbook_meta(
            source_file=source_file,
            tab="6.0  Pricing",
            location=f"Section 6.1 row {r} (US col D, PH col F, DR col G, UK col I)",
            value_label="CS Live Agent",
            calculation="Live Agent row — DR nearshore vs PH offshore per workbook",
            snapshot=snap,
        ),
    }
    for i, row in enumerate(rows):
        if "live agent" in row.get("label", "").lower():
            rows[i] = canon
            return
    rows.insert(0, canon)


def extract_cola_assumptions(ws, vendor_id: str) -> dict[str, Any]:
    """Tab 6.0 rows ~40–52: COLA % by geography and economic treatment label."""
    treatment_map = {
        "cognizant": "Bundled (inclusive)",
        "exl": "Exclusive",
        "genpact": "Exclusive",
        "ibm": 'Deferred ("mutually agreed")',
        "sutherland": "Inclusive Y1–Y5",
        "ubiquity": "Capped",
    }
    header_at: Optional[int] = None
    for r in range(35, 54):
        b = cell_str(ws.cell(r, 2).value).lower()
        if "annual cola assumption" in b:
            header_at = r
            break
    geokey = {"united states": "us", "usa": "us", "canada": "ca", "philippines": "ph", "united kingdom": "uk", "netherlands": "nl", "germany": "de"}
    out_geo: dict[str, str] = {k: "—" for k in ("us", "ca", "ph", "uk", "nl", "de")}
    ibm_note = ""
    if header_at is not None:
        r = header_at + 1
        while r <= header_at + 15:
            label = cell_str(ws.cell(r, 2).value)
            if not label or label.lower().startswith("6.3"):
                break
            key = geokey.get(label.lower())
            if not key:
                r += 1
                continue
            va = ws.cell(r, 3).value
            if vendor_id == "ibm":
                com = cell_str(ws.cell(r, 4).value)
                if key == "us" and com:
                    ibm_note = com[:240]
            if isinstance(va, (int, float)) and not (isinstance(va, float) and math.isnan(va)):
                pct = float(va) * 100 if abs(float(va)) <= 1.0 else float(va)
                if abs(pct) < 0.0005:
                    out_geo[key] = "0%"
                else:
                    out_geo[key] = f"{pct:g}%"
            else:
                ts = cell_str(va)
                if ts:
                    tl = ts.lower()
                    if "lesser" in tl or "inflation" in tl or "3%" in tl:
                        out_geo[key] = "≤3%"
                    elif "actual" in tl:
                        out_geo[key] = "actuals"
                    elif not ts and vendor_id == "cognizant" and key in ("nl", "de"):
                        out_geo[key] = "—"
                    else:
                        out_geo[key] = ts[:32]
                elif vendor_id == "cognizant" and key in ("nl", "de"):
                    out_geo[key] = "—"
                else:
                    out_geo[key] = "—"
            r += 1
    parts = [f"US {out_geo['us']}", f"CA {out_geo['ca']}", f"PH {out_geo['ph']}", f"UK {out_geo['uk']}", f"NL {out_geo['nl']}", f"DE {out_geo['de']}"]
    summary = "; ".join(parts) + f" — {treatment_map.get(vendor_id, '—')}"
    note = ibm_note if vendor_id == "ibm" else ""
    return {
        "us": out_geo["us"],
        "ca": out_geo["ca"],
        "ph": out_geo["ph"],
        "uk": out_geo["uk"],
        "nl": out_geo["nl"],
        "de": out_geo["de"],
        "treatment": treatment_map.get(vendor_id, "—"),
        "summary": summary,
        "ibmColaNote": note,
    }


def cell_str(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float) and math.isnan(v):
        return ""
    return str(v).strip()


def load_workbook_fixed(path: Path):
    """EXL file is PK zip xlsx with .XLS extension."""
    p = str(path)
    if path.name.upper().startswith("EXL_APPEND") or path.suffix.upper() == ".XLS":
        try:
            with open(path, "rb") as f:
                sig = f.read(4)
            if sig == b"PK\x03\x04":
                tmp = F1 / "_exl_converted.xlsx"
                shutil.copy(path, tmp)
                return openpyxl.load_workbook(tmp, data_only=True)
        except OSError:
            pass
    return openpyxl.load_workbook(p, data_only=True)


def find_total_annual_cost_row(ws, start_row: int = 100) -> Optional[int]:
    for r in range(start_row, min(ws.max_row or 200, 250) + 1):
        b = cell_str(ws.cell(r, 2).value).lower()
        if "total annual cost" in b:
            return r
    return None


def quarterly_values(ws, row: int, last_col: int = 22) -> list[Optional[float]]:
    """Columns C (3) through last_col inclusive (default V=22 → 20 quarters)."""
    out: list[Optional[float]] = []
    for c in range(3, last_col + 1):
        v = ws.cell(row, c).value
        if v is None or (isinstance(v, float) and math.isnan(v)):
            out.append(None)
        elif isinstance(v, (int, float)):
            out.append(float(v))
        else:
            out.append(None)
    return out


def quarterly_values_extend(ws, row: int, max_quarters: int = 32) -> list[Optional[float]]:
    q: list[Optional[float]] = []
    c = 3
    while c <= 3 + max_quarters - 1:
        v = ws.cell(row, c).value
        if v is None or (isinstance(v, float) and math.isnan(v)):
            q.append(None)
        elif isinstance(v, (int, float)):
            q.append(float(v))
        else:
            q.append(None)
        c += 1
    while q and q[-1] is None:
        q.pop()
    return q


def infer_scale_million(first_nonzero: float) -> bool:
    """True if values are raw dollars and should divide by 1e6."""
    return abs(first_nonzero) > 1000


def annualize_quarters(
    q: list[Optional[float]], num_years: int = 5
) -> tuple[list[Optional[float]], float]:
    """Sum quarters into annual buckets; TCV = sum of all non-None quarters."""
    years: list[Optional[float]] = []
    for yi in range(num_years):
        chunk = q[yi * 4 : (yi + 1) * 4]
        if all(x is None for x in chunk):
            years.append(None)
        else:
            s = sum(x for x in chunk if x is not None)
            years.append(s)
    tcv = sum(x for x in q if x is not None)
    return years, tcv


def scale_years_tcv(
    years: list[Optional[float]], tcv: float, divide_million: bool
) -> tuple[list[Optional[float]], float]:
    if not divide_million:
        return years, tcv
    ny = [y / 1e6 if y is not None else None for y in years]
    return ny, tcv / 1e6


def parse_ubiquity_annual_from_rows(ws, r110: int = 110, r111: int = 111) -> Optional[tuple[list[Optional[float]], float]]:
    q: list[Optional[float]] = []
    for c in range(3, 23):
        a = ws.cell(r110, c).value
        b = ws.cell(r111, c).value
        sa = float(a) if isinstance(a, (int, float)) and not (isinstance(a, float) and math.isnan(a)) else 0.0
        sb = float(b) if isinstance(b, (int, float)) and not (isinstance(b, float) and math.isnan(b)) else 0.0
        tot = sa + sb if (sa or sb) else None
        q.append(tot)
    years, tcv = annualize_quarters(q)
    years2, tcv2 = scale_years_tcv(years, tcv, divide_million=True)
    return years2, tcv2


def parse_ubiquity_row112(ws, row: int = 112) -> Optional[tuple[list[Optional[float]], float]]:
    q = quarterly_values(ws, row)
    years, tcv = annualize_quarters(q)
    nz = next((x for x in q if x is not None), None)
    if nz is None:
        return None
    div = infer_scale_million(nz)
    years2, tcv2 = scale_years_tcv(years, tcv, divide_million=div)
    return years2, tcv2


def count_governance(ws, r0: int = 33, r1: int = 44) -> dict[str, int]:
    """Tab 5.0 governance rows ~33–44; vendor stance in column C only (avoid double-count from detail cols)."""
    commit = partial = cannot = 0
    for r in range(r0, r1 + 1):
        req = cell_str(ws.cell(r, 2).value).lower()
        if not req or req.startswith("requirement"):
            continue
        t = cell_str(ws.cell(r, 3).value).lower()
        if not t:
            continue
        if "cannot commit" in t or t.startswith("cannot") or "unable" in t:
            cannot += 1
        elif "partial" in t:
            partial += 1
        elif "commit" in t or t in ("y", "yes"):
            commit += 1
    return {"commit": commit, "partial": partial, "cannotCommit": cannot}


def extract_governance_items(ws, r0: int = 33, r1: int = 44) -> list[dict[str, Any]]:
    """Line-level Tab 5.0 commitments for dashboard governance drill-down."""
    out: list[dict[str, Any]] = []
    for r in range(r0, r1 + 1):
        req = cell_str(ws.cell(r, 2).value)
        if not req or req.lower().startswith("requirement"):
            continue
        t = cell_str(ws.cell(r, 3).value).lower()
        if not t:
            continue
        if "cannot commit" in t or t.startswith("cannot") or "unable" in t:
            st = "cannot"
        elif "partial" in t:
            st = "partial"
        elif "commit" in t or t in ("y", "yes"):
            st = "commit"
        else:
            continue
        label = req.strip()
        if len(label) > 220:
            label = label[:217] + "…"
        out.append({"status": st, "label": label})
    return out


def onetime_source(ws, row: int, label: str, source_file: str) -> dict[str, Any]:
    snap = snapshot_around_row(ws, row)
    return workbook_meta(
        source_file=source_file,
        tab="6.0  Pricing",
        location=f"Row {row} (one-time / transition line)",
        value_label=label,
        calculation="Sum of populated quarterly cells across C:V for this row",
        snapshot=snap,
    )


def extract_efficiency_table(ws, source_file: str = "", tab_label: str = "9.0 Client EfficiencyAssumption") -> dict[str, Any]:
    """Read Tab 9.0: row 6 headers from col C; rows 8–12 data (col B = geography)."""
    from openpyxl.utils import get_column_letter

    headers: list[str] = []
    for c in range(3, 14):
        h = cell_str(ws.cell(6, c).value)
        if not h:
            break
        headers.append(h)
    rows_out: list[dict[str, Any]] = []
    for r in range(8, 13):
        label = cell_str(ws.cell(r, 2).value)
        if not label:
            continue
        cells: dict[str, str] = {}
        cell_provenance: dict[str, Any] = {}
        for i, h in enumerate(headers):
            col = 3 + i
            v = ws.cell(r, col).value
            if v is None or (isinstance(v, float) and math.isnan(v)):
                cells[h] = ""
            elif isinstance(v, float) and "%" in h and 0 <= abs(v) <= 1:
                cells[h] = f"{v * 100:.2f}%"
            else:
                cells[h] = cell_str(v)
            if source_file and cells[h]:
                L = get_column_letter(col)
                ref = f"{L}{r}"
                cell_provenance[h] = workbook_meta(
                    source_file=source_file,
                    tab=tab_label,
                    location=f"Cell {ref} · {label} · {h}",
                    value_label=cells[h][:120],
                    calculation="Tab 9.0 vendor-reported cell",
                    snapshot=cell_snapshot(ws, ref),
                )
        row_obj: dict[str, Any] = {"geography": label, "cells": cells}
        if cell_provenance:
            row_obj["cellProvenance"] = cell_provenance
        rows_out.append(row_obj)
    return {"headers": headers, "rows": rows_out}


def sum_onetime_row(ws, row: int) -> Optional[float]:
    total = 0.0
    anyv = False
    for c in range(3, 23):
        v = ws.cell(row, c).value
        if isinstance(v, (int, float)) and not (isinstance(v, float) and math.isnan(v)):
            total += float(v)
            anyv = True
    return total if anyv else None


def docx_paragraph_texts(path: Path) -> str:
    try:
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


def tab_status(wb: openpyxl.Workbook, required: list[str]) -> list[dict[str, str]]:
    names = set(wb.sheetnames)
    out = []
    for t in required:
        if t in names:
            ws = wb[t]
            nonempty = 0
            for row in ws.iter_rows(max_row=min(ws.max_row or 0, 200), max_col=12):
                for cell in row:
                    if cell.value not in (None, ""):
                        nonempty += 1
            if nonempty < 8:
                status = "partial"
            else:
                status = "complete"
            out.append({"tab": t, "status": status})
        else:
            out.append({"tab": t, "status": "missing"})
    return out


STANDARD_TABS = [
    "1.0  Table of Contents",
    "2.0  Industry Experience",
    "3.0  Relevant Experience",
    "4.0  Workforce & Delivery",
    "5.0  Technology Solution",
    "6.0  Pricing",
    "7.0  Client Migration",
    "8.0  Regulatory Compliance",
    "9.0 Client EfficiencyAssumption",
]

DRILL_TABS = [
    "2.0  Industry Experience",
    "3.0  Relevant Experience",
    "4.0  Workforce & Delivery",
    "5.0  Technology Solution",
    "7.0  Client Migration",
    "8.0  Regulatory Compliance",
]

MIN_QUESTION_WORDS = 8
MIN_RESPONSE_CHARS = 2

INSTRUCTION_PATTERNS = (
    "complete column d",
    "for rows asking for years of experience",
    "enter the number",
    "provide a written response",
    "specific, evidence-based answers",
    "what we are asking",
    "vendor response",
    "evidence-based answers are expected",
)

INSTRUCTION_SKIP = (
    "complete column d for every row",
    "for rows asking for years of experience",
    "enter the number of years",
    "provide a written response",
    "specific, evidence-based answers are expected",
    "what we are asking",
)


def is_instruction_text(text: str) -> bool:
    lower = (text or "").strip().lower()
    return any(pat in lower for pat in INSTRUCTION_SKIP)


def is_template_or_instruction(text: str) -> bool:
    lower = (text or "").strip().lower()
    if not lower:
        return True
    if lower in ("no response provided.", "no response provided"):
        return True
    for pat in INSTRUCTION_PATTERNS:
        if pat in lower:
            # "vendor response" is a section label inside large merged cells; do not flag long prose answers.
            if pat == "vendor response" and len(lower) > 400:
                continue
            return True
    return len(lower) < 3


def is_merged_layout_cell(text: str) -> bool:
    if not text or len(text) < 400:
        return False
    lower = text.lower().strip()
    if lower.startswith("what we are asking"):
        return True
    markers = ("describe your", "confirm:", "how many years", "provide at least", "what is your", "explain your", "list your")
    hits = sum(1 for m in markers if m in lower)
    if hits >= 2 and len(text) > 1500:
        return True
    return False


def split_merged_qa(text: str) -> tuple[str, str]:
    parts = text.split("\n\n")
    if len(parts) >= 2:
        potential_q = parts[0].strip()
        potential_a = "\n\n".join(parts[1:]).strip()
        if len(potential_q) < 500 and len(potential_a) > 50:
            return potential_q, potential_a
    return "", text


def filter_drill_instruction_snippets(snippets: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Skip instruction-only answer cells; strip instruction-only question headers when answer is real."""
    out: list[dict[str, Any]] = []
    for s in snippets:
        t = (s.get("text") or "").strip()
        q = (s.get("questionText") or "").strip()
        if is_instruction_text(t):
            continue
        s2 = dict(s)
        if is_instruction_text(q) and t and not is_instruction_text(t):
            s2["questionText"] = ""
        out.append(s2)
    return out


def dedup_drill_snippets(snippets: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Remove snippets whose text is a substring of a longer kept snippet; re-attach unanswered 'no response' rows."""
    snippets = filter_drill_instruction_snippets(snippets)
    if len(snippets) <= 1:
        return snippets

    texts = [(i, (s.get("text") or "").strip()) for i, s in enumerate(snippets)]
    texts.sort(key=lambda x: -len(x[1]))
    kept: list[dict[str, Any]] = []
    kept_texts: list[str] = []
    for idx, txt in texts:
        if not txt or txt.lower() in ("no response provided.", "no response provided"):
            continue
        if is_merged_layout_cell(txt):
            continue
        if any(txt in kt for kt in kept_texts):
            continue
        kept.append(snippets[idx])
        kept_texts.append(txt)

    answered_qs = {(s.get("questionText") or "").strip()[:220] for s in kept}
    for s in snippets:
        tl = (s.get("text") or "").strip().lower()
        if tl not in ("no response provided.", "no response provided", ""):
            continue
        q = (s.get("questionText") or "").strip()
        qn = q[:220]
        if not q or qn in answered_qs or is_instruction_text(q) or is_template_or_instruction(q):
            continue
        kept.append(s)
        answered_qs.add(qn)
    return kept


def looks_like_questionnaire_question(t: str) -> bool:
    s = (t or "").strip()
    if len(s) < 10:
        return False
    if "?" in s:
        return True
    low = s.lower()
    keys = (
        "describe ",
        "confirm:",
        "how many",
        "provide ",
        "list ",
        "what is your",
        "explain your",
        "state whether",
        "indicate if",
        "are you ",
        "will you",
        "does your",
    )
    return any(k in low for k in keys)


def looks_like_prose_answer(t: str) -> bool:
    s = (t or "").strip()
    if len(s) < 50:
        return False
    if looks_like_questionnaire_question(s) and len(s) < 120:
        return False
    if s.count("?") == 1 and len(s) < 180:
        return False
    return True


def normalize_for_consensus(s: str) -> str:
    return " ".join(s.lower().split())


@dataclass
class DrillContext:
    question_keys: set[tuple[str, int, int]]
    canonical_norm: dict[tuple[str, int, int], str]
    canonical_raw: dict[tuple[str, int, int], str]
    template_norms: set[str]


def build_drill_consensus(wbs: dict[str, Optional[openpyxl.Workbook]]) -> tuple[set[tuple[str, int, int]], dict[tuple[str, int, int], str], dict[tuple[str, int, int], str]]:
    """Cells where 3+ vendors share identical long text = RFP template question cells."""
    bucket: dict[tuple[str, int, int], list[tuple[str, str, str]]] = defaultdict(list)
    for vid, wb in wbs.items():
        if wb is None:
            continue
        for tab in DRILL_TABS:
            if tab not in wb.sheetnames:
                continue
            ws = wb[tab]
            max_r = min(ws.max_row or 280, 320)
            for r in range(1, max_r + 1):
                for c in range(2, 8):
                    raw = cell_str(ws.cell(r, c).value)
                    if len(raw.split()) < MIN_QUESTION_WORDS:
                        continue
                    nt = normalize_for_consensus(raw)
                    bucket[(tab, r, c)].append((vid, nt, raw))
    question_keys: set[tuple[str, int, int]] = set()
    canonical_norm: dict[tuple[str, int, int], str] = {}
    canonical_raw: dict[tuple[str, int, int], str] = {}
    for key, pairs in bucket.items():
        by_norm: dict[str, list[tuple[str, str]]] = defaultdict(list)
        for vid, nt, raw in pairs:
            by_norm[nt].append((vid, raw))
        for nt, lst in by_norm.items():
            if len({v for v, _ in lst}) >= 3:
                raw_pick = min((raw for _, raw in lst), key=len)
                question_keys.add(key)
                canonical_norm[key] = nt
                canonical_raw[key] = raw_pick
                break
    return question_keys, canonical_norm, canonical_raw


def load_template_norm_strings() -> set[str]:
    norms: set[str] = set()
    if not F1.is_dir():
        return norms
    candidates = []
    for pat in ("*0320*.xlsx", "*Vendor Response Workbook*.xlsx", "*Appendix B*TMS*.xlsx"):
        candidates.extend([p for p in F1.glob(pat) if "Cognizant" not in p.name and "Genpact" not in p.name and "IBM" not in p.name][:8])
    seen = set()
    for p in candidates[:5]:
        if p.name in seen:
            continue
        seen.add(p.name)
        try:
            twb = openpyxl.load_workbook(p, data_only=True, read_only=True)
            for tab in DRILL_TABS:
                if tab not in twb.sheetnames:
                    continue
                ws = twb[tab]
                for row in ws.iter_rows(min_row=1, max_row=320, min_col=2, max_col=7, values_only=True):
                    for val in row:
                        t = cell_str(val)
                        if len(t.split()) >= MIN_QUESTION_WORDS:
                            norms.add(normalize_for_consensus(t))
            twb.close()
        except Exception:
            continue
    return norms


def _is_tab7_migration_template_row_text(t: str) -> bool:
    """True when the cell is FIS RFP instructions, not vendor wave narrative."""
    if len(t) < 100:
        return False
    low = t.lower()
    if "de-identified client dataset will be provided" in low[:1200]:
        return True
    if "using the supplemental data provided, please architect waves" in low:
        return True
    if "define a raci for the client migration" in low and len(t) < 900:
        return True
    return False


def extract_migration_notes_7(wb: openpyxl.Workbook) -> str:
    tab = "7.0  Client Migration"
    if tab not in wb.sheetnames:
        return ""
    ws = wb[tab]
    chunks: list[str] = []
    max_r = min(ws.max_row or 0, 420)
    for r in range(1, max_r + 1):
        best = ""
        for c in range(2, 16):
            t = cell_str(ws.cell(r, c).value)
            if len(t.split()) >= 6 and len(t) > len(best):
                best = t
        if best:
            chunks.append(best)
    kept = [c for c in chunks if not _is_tab7_migration_template_row_text(c)]
    use_chunks = kept if kept else chunks
    # Tab 7.0 tables repeat the same column header (e.g. service line) on every row; drop exact duplicates.
    deduped: list[str] = []
    seen_norm: set[str] = set()
    for c in use_chunks[:48]:
        k = normalize_for_consensus(c)
        if not k or k in seen_norm:
            continue
        seen_norm.add(k)
        deduped.append(c)
    merged = "\n\n".join(deduped[:24])
    return merged[:8000] if merged else ""


def sanitize_migration_notes_text(text: str) -> str:
    """Drop leading RFP boilerplate and repeated Part 2 instruction block from Tab 7.0 merge."""
    if not text or text.strip() == "—":
        return text
    t = text.strip()
    low = t.lower()
    for anchor in ("this table represent", "this table represents"):
        idx = low.find(anchor)
        if idx > 80:
            t = t[idx:].lstrip()
            low = t.lower()
            break
    part2_boiler = re.compile(
        r"\n{2,}Part\s+2\s*[\u2014\u2013\-–]\s*Wave Timeline:.+?"
        r"evaluated more favorably than generic frameworks\.?\s*",
        re.IGNORECASE | re.DOTALL,
    )
    t = part2_boiler.sub("\n\n", t)
    return t.strip() or text.strip()


def extract_migration_wave_count_hint(wb: openpyxl.Workbook) -> str:
    """Append max 'Wave N' index from Tab 7.0 so migration text mentions wave counts (validates analyst checks)."""
    tab = "7.0  Client Migration"
    if tab not in wb.sheetnames:
        return ""
    ws = wb[tab]
    nums: list[int] = []
    max_r = min(ws.max_row or 0, 520)
    max_c = min(ws.max_column or 8, 14)
    for r in range(1, max_r + 1):
        for c in range(1, max_c + 1):
            v = ws.cell(r, c).value
            if not isinstance(v, str):
                continue
            for m in re.finditer(r"(?i)\bwave\s*(\d{1,2})\b", v):
                nums.append(int(m.group(1)))
    if not nums or max(nums) < 4:
        return ""
    return f"Vendor documents up to {max(nums)} migration waves in Tab 7.0 (enumerated wave list / summary)."


def migration_notes_with_waves(wb: openpyxl.Workbook) -> str:
    base = sanitize_migration_notes_text(extract_migration_notes_7(wb) or "—")
    hint = extract_migration_wave_count_hint(wb)
    if not hint:
        return base
    n = hint.split("up to ", 1)[1].split(" ", 1)[0] if "up to " in hint else ""
    if n.isdigit() and re.search(rf"(?i)\b{n}\b.*\bwaves?\b", base):
        return base
    return f"{base}\n\n{hint}"


def _duplicate_question_answer(response: str, q_raw: str, q_label: str, ctx: DrillContext) -> bool:
    """True if `response` is still template/question text, not a vendor answer."""
    rt = normalize_for_consensus(response)
    if not rt.strip():
        return True
    if rt == normalize_for_consensus(q_raw) or rt == normalize_for_consensus(q_label):
        return True
    if len(response.split()) >= MIN_QUESTION_WORDS and rt in ctx.template_norms:
        return True
    return False


def _merge_tabular_row_cells(ws, r: int, min_each: int = 8, last_col: int = 14) -> str:
    """Join non-empty cells from column C onward for table-style answers (EXL / multi-column layouts)."""
    parts: list[str] = []
    xc = min(last_col, ws.max_column or last_col)
    for c in range(3, xc + 1):
        t = cell_str(ws.cell(r, c).value)
        if len(t) >= min_each:
            parts.append(t.strip())
    # Avoid repeating the same cell text when multiple columns hold the same label (common in wave tables).
    deduped: list[str] = []
    prev_norm = ""
    for t in parts:
        n = normalize_for_consensus(t)
        if n == prev_norm:
            continue
        deduped.append(t)
        prev_norm = n
    return "\n\n".join(deduped)


def _likely_column_header_row(ws, r: int) -> bool:
    """True if row looks like a header row (short labels in every populated cell)."""
    parts: list[str] = []
    for c in range(3, min(12, (ws.max_column or 12) + 1)):
        t = cell_str(ws.cell(r, c).value)
        if len(t) > 3:
            parts.append(t)
    if len(parts) < 3:
        return False
    return all(len(p) < 72 for p in parts)


def _collect_tabular_answer_below(
    ws,
    tab: str,
    ctx: DrillContext,
    question_row: int,
    q_raw: str,
    q_label: str,
    max_down: int = 26,
) -> tuple[str, str]:
    """When the question row has no same-row answer, merge substantive cells from following rows until the next template question row."""
    parts: list[str] = []
    first_ref = ""
    sheet_hi = ws.max_row or question_row
    end_row = min(question_row + max_down, sheet_hi)
    for rr in range(question_row + 1, end_row + 1):
        sub_q = False
        for c in range(2, 8):
            raw = cell_str(ws.cell(rr, c).value)
            if len(raw.split()) < max(4, MIN_QUESTION_WORDS - 2):
                continue
            if (tab, rr, c) in ctx.question_keys or normalize_for_consensus(raw) in ctx.template_norms:
                sub_q = True
                break
        merged = _merge_tabular_row_cells(ws, rr, min_each=6)
        if sub_q and len(merged) < 45:
            break
        if _likely_column_header_row(ws, rr):
            continue
        if len(merged) < 25:
            continue
        if _duplicate_question_answer(merged, q_raw, q_label, ctx):
            continue
        if not first_ref:
            first_ref = ws.cell(rr, 3).coordinate
            t3 = cell_str(ws.cell(rr, 3).value)
            if len(t3) < 8:
                for c in range(4, min(10, (ws.max_column or 9) + 1)):
                    if len(cell_str(ws.cell(rr, c).value)) >= 8:
                        first_ref = ws.cell(rr, c).coordinate
                        break
        parts.append(merged)
    text = "\n\n".join(parts)[:4500]
    return text, first_ref


def extract_drill_tab(wb: openpyxl.Workbook, tab: str, ctx: DrillContext, max_rows: int = 48) -> dict[str, Any]:
    if tab not in wb.sheetnames:
        return {"tab": tab, "missing": True, "snippets": []}
    ws = wb[tab]
    snippets: list[dict[str, Any]] = []
    seen_rows: set[int] = set()
    max_r = min(ws.max_row or 280, 360)
    section_title = tab
    for r in range(1, max_r + 1):
        if len(snippets) >= max_rows:
            break
        b2 = cell_str(ws.cell(r, 2).value)
        if re.match(r"^\d+\.\d+\s+", b2) and len(b2) < 220:
            section_title = b2.strip()[:500]

        q_candidates: list[tuple[int, str, str]] = []
        for c in range(2, 8):
            raw = cell_str(ws.cell(r, c).value)
            if len(raw.split()) < max(4, MIN_QUESTION_WORDS - 2):
                continue
            key_gc = (tab, r, c)
            if key_gc in ctx.question_keys:
                q_candidates.append((c, raw, ctx.canonical_raw.get(key_gc, raw)))
            elif normalize_for_consensus(raw) in ctx.template_norms:
                q_candidates.append((c, raw, raw))

        if not q_candidates:
            merged_early = _merge_tabular_row_cells(ws, r, min_each=6)
            other_populated = sum(1 for c in range(3, min(13, (ws.max_column or 12) + 1)) if len(cell_str(ws.cell(r, c).value)) > 4)
            short_lead = len(b2) <= 6 or b2.strip() in ("#", "—", "-") or (b2.isdigit() and len(b2) <= 3)
            label_row = 8 <= len(b2) <= 90 and other_populated >= 3
            if (
                len(merged_early) >= 40
                and other_populated >= 2
                and (short_lead or label_row)
                and not _likely_column_header_row(ws, r)
                and r not in seen_rows
            ):
                if not _duplicate_question_answer(merged_early, b2, section_title, ctx) and not is_merged_layout_cell(merged_early):
                    seen_rows.add(r)
                    snippets.append(
                        {
                            "ref": ws.cell(r, 3).coordinate,
                            "questionText": section_title[:1200],
                            "text": merged_early[:4500],
                        }
                    )
            continue

        q_col, q_raw, tpl_raw = min(q_candidates, key=lambda x: x[0])
        key = (tab, r, q_col)
        if not tpl_raw and normalize_for_consensus(q_raw) in ctx.template_norms:
            tpl_raw = q_raw
        canonical = ctx.canonical_raw.get(key, "")
        q_label = (canonical or tpl_raw or q_raw).strip()
        if len(q_label) > 900:
            q_label = q_label[:900] + "…"

        response_text = ""
        resp_ref = ""
        pref = (tpl_raw or canonical or "").strip()
        qr = q_raw.strip()
        if pref and qr.startswith(pref) and len(qr) > len(pref) + MIN_RESPONSE_CHARS:
            response_text = qr[len(pref) :].strip()
            resp_ref = ws.cell(r, q_col).coordinate
        elif canonical and qr.startswith(canonical.strip()) and len(qr) > len(canonical.strip()) + MIN_RESPONSE_CHARS:
            response_text = qr[len(canonical.strip()) :].strip()
            resp_ref = ws.cell(r, q_col).coordinate

        if not response_text or _duplicate_question_answer(response_text, q_raw, q_label, ctx):
            candidates: list[tuple[int, str, str]] = []
            for c in range(2, min(15, (ws.max_column or 14) + 1)):
                if c == q_col:
                    continue
                raw = cell_str(ws.cell(r, c).value)
                if not raw.strip():
                    continue
                nk = (tab, r, c)
                if nk in ctx.question_keys:
                    continue
                if normalize_for_consensus(raw) == normalize_for_consensus(q_raw):
                    continue
                if _duplicate_question_answer(raw, q_raw, q_label, ctx):
                    continue
                if is_merged_layout_cell(raw):
                    continue
                if is_template_or_instruction(raw) and len(raw) < 120:
                    continue
                candidates.append((len(raw), raw, ws.cell(r, c).coordinate))
            if candidates:
                candidates.sort(key=lambda x: -x[0])
                response_text = candidates[0][1]
                resp_ref = candidates[0][2]

        if not response_text or _duplicate_question_answer(response_text, q_raw, q_label, ctx):
            merged = _merge_tabular_row_cells(ws, r, min_each=10)
            if is_merged_layout_cell(merged):
                merged = ""
            if (
                len(merged) > 50
                and not _likely_column_header_row(ws, r)
                and not _duplicate_question_answer(merged, q_raw, q_label, ctx)
            ):
                response_text = merged
                resp_ref = ws.cell(r, 3).coordinate

        if response_text and is_merged_layout_cell(response_text):
            split_q, split_a = split_merged_qa(response_text)
            if split_a and len(split_a) > 80 and not is_merged_layout_cell(split_a):
                response_text = split_a
                if split_q and ("?" in split_q or looks_like_questionnaire_question(split_q)) and not is_template_or_instruction(split_q):
                    q_label = split_q[:900] + ("…" if len(split_q) > 900 else "")
            else:
                response_text = ""

        if not response_text or _duplicate_question_answer(response_text, q_raw, q_label, ctx):
            below, ref_below = _collect_tabular_answer_below(ws, tab, ctx, r, q_raw, q_label)
            if len(below) >= 40 and not is_merged_layout_cell(below):
                response_text = below
                resp_ref = ref_below or resp_ref

        if not response_text or _duplicate_question_answer(response_text, q_raw, q_label, ctx):
            for dc in (-1, 1):
                cc = q_col + dc
                if cc < 2:
                    continue
                mx = ws.max_column or 14
                if cc > mx:
                    continue
                adj = cell_str(ws.cell(r, cc).value).strip()
                if len(adj) < 28:
                    continue
                if (tab, r, cc) in ctx.question_keys:
                    continue
                if is_merged_layout_cell(adj) or is_template_or_instruction(adj):
                    continue
                if _duplicate_question_answer(adj, q_raw, q_label, ctx):
                    continue
                response_text = adj[:4500]
                resp_ref = ws.cell(r, cc).coordinate
                break

        if not response_text or _duplicate_question_answer(response_text, q_raw, q_label, ctx):
            long_q = q_raw if len(q_raw) > 200 else ""
            if long_q:
                sq, sa = split_merged_qa(long_q)
                if sa and len(sa) > 50 and sq and len(sq) < 480:
                    response_text = sa[:4500]
                    resp_ref = ws.cell(r, q_col).coordinate
                    if not is_template_or_instruction(sq):
                        q_label = sq[:900] + ("…" if len(sq) > 900 else "")

        if not response_text or _duplicate_question_answer(response_text, q_raw, q_label, ctx):
            response_text = "No response provided."
            resp_ref = resp_ref or ws.cell(r, q_col).coordinate

        if r in seen_rows:
            continue
        seen_rows.add(r)
        snippets.append(
            {
                "ref": resp_ref,
                "questionText": q_label[:1200],
                "text": response_text[:4500],
            }
        )
    snippets = dedup_drill_snippets(snippets)
    return {"tab": tab, "missing": False, "snippets": snippets}


def extract_drill_blocks(wb: openpyxl.Workbook, ctx: DrillContext, max_per_tab: int = 48) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for tab in DRILL_TABS:
        cap = 130 if "7.0" in tab else max_per_tab
        out.append(extract_drill_tab(wb, tab, ctx, max_rows=cap))
    return out


def extract_questionnaire_blocks(qwb: openpyxl.Workbook, source_file: str, max_per_sheet: int = 6) -> list[dict[str, Any]]:
    """Ubiquity RFP questionnaire: B/C/D Q&A with column-order fixes when prose lands in the wrong cell."""
    blocks: list[dict[str, Any]] = []
    for sname in qwb.sheetnames[:8]:
        ws = qwb[sname]
        sn: list[dict[str, Any]] = []
        max_r = min(ws.max_row or 120, 200)
        for r in range(1, max_r + 1):
            b = cell_str(ws.cell(r, 2).value).strip()
            c = cell_str(ws.cell(r, 3).value).strip()
            d = cell_str(ws.cell(r, 4).value).strip()
            qtext, rtext, ref = "", "", ws.cell(r, 3).coordinate

            b_q, c_q, d_q = looks_like_questionnaire_question(b), looks_like_questionnaire_question(c), looks_like_questionnaire_question(d)
            b_a, c_a = looks_like_prose_answer(b), looks_like_prose_answer(c)

            if b_q and len(b.split()) >= 6:
                qtext = b
                rtext = d or c
                ref = ws.cell(r, 4).coordinate if d else ws.cell(r, 3).coordinate
            elif c_q and len(c.split()) >= 6:
                qtext = c
                rtext = d or b
                ref = ws.cell(r, 4).coordinate if d else ws.cell(r, 2).coordinate
            elif d_q and len(d.split()) >= 6:
                qtext = d
                rtext = c or b
                ref = ws.cell(r, 3).coordinate if c else ws.cell(r, 2).coordinate
            elif b_a and len(b.split()) >= 12 and not b_q:
                rtext = b
                qtext = c if c_q else (d if d_q else "")
                ref = ws.cell(r, 2).coordinate
                if not qtext.strip() and c_q:
                    qtext = c
                    rtext = d or b
                    ref = ws.cell(r, 3).coordinate

            if not rtext.strip() and b_a and not b_q and (c_q or (c and "?" in c)):
                qtext, rtext = c, b
                ref = ws.cell(r, 2).coordinate

            if qtext and not rtext.strip() and looks_like_prose_answer(qtext) and "?" not in qtext and len(qtext) > 80:
                rtext, qtext = qtext, ""

            if not qtext.strip() and rtext.strip():
                qtext = "Vendor response"

            if not qtext.strip():
                continue
            if is_template_or_instruction(qtext):
                continue

            if not rtext.strip():
                for cc in range(2, min(12, (ws.max_column or 11) + 1)):
                    if cc == 2 and b:
                        continue
                    t = cell_str(ws.cell(r, cc).value).strip()
                    if len(t) < 30 or is_template_or_instruction(t):
                        continue
                    if looks_like_questionnaire_question(t) and len(t) < 400:
                        continue
                    rtext = t
                    ref = ws.cell(r, cc).coordinate
                    break

            if not rtext.strip():
                rtext = "No response provided."
            if is_merged_layout_cell(rtext):
                continue

            sn.append({"ref": ref or f"C{r}", "questionText": qtext[:1200], "text": rtext[:2400]})
            if len(sn) >= max_per_sheet:
                break
        sn = dedup_drill_snippets(sn)
        blocks.append({"tab": f"Q: {sname}", "missing": False, "snippets": sn})
    return blocks if blocks else [{"tab": t, "missing": True, "snippets": []} for t in DRILL_TABS]


# Legacy name kept for ubiquity fallback paths
def extract_drill_snippets(wb: openpyxl.Workbook, ctx: Optional[DrillContext] = None, max_per_tab: int = 48) -> list[dict[str, Any]]:
    if ctx is not None:
        return extract_drill_blocks(wb, ctx, max_per_tab=max_per_tab)
    empty = DrillContext(set(), {}, {}, set())
    return extract_drill_blocks(wb, empty, max_per_tab=max_per_tab)


@dataclass
class VendorExtract:
    id: str
    displayName: str
    color: str
    workbookPath: str
    pricing: dict[str, Any] = field(default_factory=dict)
    rateCard: list = field(default_factory=list)
    colaAssumptions: dict[str, Any] = field(default_factory=dict)
    efficiency: dict[str, Any] = field(default_factory=dict)
    governance: dict[str, int] = field(default_factory=dict)
    governanceItems: list = field(default_factory=list)
    oneTimeLines: list = field(default_factory=list)
    migrationNotes: str = ""
    flags: list = field(default_factory=list)
    evaluationBriefText: str = ""
    adminTabs: list = field(default_factory=list)
    sources: dict[str, str] = field(default_factory=dict)
    drilldownSnippets: list = field(default_factory=list)


def extract_cognizant(wb: Optional[openpyxl.Workbook] = None, drill_ctx: Optional[DrillContext] = None) -> VendorExtract:
    path = F1 / "Cognizant_ Appendix B - TMS RFP - Cognizant - Vendor Response Workbook.xlsx"
    if wb is None:
        wb = openpyxl.load_workbook(path, data_only=True)
    if drill_ctx is None:
        drill_ctx = DrillContext(set(), {}, {}, set())
    ws = wb["6.0  Pricing"]
    row = find_total_annual_cost_row(ws) or 120
    q = quarterly_values(ws, row)
    nz = next((x for x in q if x is not None), 0)
    years, tcv = annualize_quarters(q)
    years, tcv = scale_years_tcv(years, tcv, divide_million=infer_scale_million(nz))
    brief = F4 / "Cognizant_Evaluation_Brief.docx"
    v = VendorExtract(
        id="cognizant",
        displayName="Cognizant",
        color="#1E3A5F",
        workbookPath=str(path.relative_to(BASE)),
        evaluationBriefText=docx_paragraph_texts(brief),
        flags=[
            "Proposal cites ~$25M vendor investment; workbook one-time tech lines (~$5.2M excl. absorbed training) materially lower — pressure-test reconciliation.",
            "Efficiency certainty shown as 100%* in Tab 9.0; asterisk conditions not specified in workbook.",
        ],
        sources={
            "totalAnnualCost": f"'6.0  Pricing'!B{row}: quarterly C{row}:V{row} (raw $, /1e6)",
            "efficiency": "'9.0 Client EfficiencyAssumption'!C8:G12",
        },
    )
    v.pricing = {
        "totalAnnualCostRow": row,
        "quarterlyUnit": "USD raw → $M",
        "years": [{"year": i + 1, "valueM": years[i]} for i in range(5)],
        "tcvM": round(tcv, 2),
        "y6M": None,
        "y7M": None,
        "tcv7M": None,
        "provenance": build_pricing_provenance(
            path.name, ws, row, round(tcv, 2), years, "raw USD → $M (quarterly sums ÷ 1e6)"
        ),
    }
    v.rateCard = rate_rows_with_snapshots(ws, path.name)
    _patch_cognizant_cs_live(v.rateCard, ws, path.name)
    v.colaAssumptions = extract_cola_assumptions(ws, "cognizant")
    eff_ws = wb["9.0 Client EfficiencyAssumption"]
    v.efficiency = extract_efficiency_table(eff_ws, path.name, "9.0 Client EfficiencyAssumption")
    _esnap = efficiency_snapshot(eff_ws)
    v.efficiency["tableSnapshot"] = _esnap
    v.efficiency["sourcePreview"] = workbook_meta(
        source_file=path.name,
        tab="9.0 Client EfficiencyAssumption",
        location="Rows 6–12 (data), row 6 headers",
        value_label="Efficiency assumptions",
        calculation="Vendor Tab 9.0 structured fields",
        snapshot=_esnap,
    )
    v.governance = count_governance(wb["5.0  Technology Solution"])
    v.governanceItems = extract_governance_items(wb["5.0  Technology Solution"])
    v.sources["governance"] = "Tab 5.0 — Commit / Partial / Cannot from vendor column C"
    v.oneTimeLines = [
        {
            "label": "Tech overlay (ex.)",
            "row": 74,
            "sumQuarterlyUsd": sum_onetime_row(ws, 74),
            "sourcePreview": onetime_source(ws, 74, "Tech overlay", path.name),
        },
        {
            "label": "Tech migration (ex.)",
            "row": 76,
            "sumQuarterlyUsd": sum_onetime_row(ws, 76),
            "sourcePreview": onetime_source(ws, 76, "Tech migration", path.name),
        },
    ]
    v.migrationNotes = migration_notes_with_waves(wb)
    v.adminTabs = [{"tab": t, "status": "complete" if t in wb.sheetnames else "missing"} for t in STANDARD_TABS]
    v.drilldownSnippets = extract_drill_blocks(wb, drill_ctx)
    enrich_drill_snippets(wb, v.drilldownSnippets, path.name)
    enrich_drill_linked_documents("cognizant", v.drilldownSnippets)
    return v


def extract_exl(wb: Optional[openpyxl.Workbook] = None, drill_ctx: Optional[DrillContext] = None) -> VendorExtract:
    path = F1 / "EXL_APPEND~1.XLS"
    if wb is None:
        wb = load_workbook_fixed(path)
    if drill_ctx is None:
        drill_ctx = DrillContext(set(), {}, {}, set())
    ws = wb["6.0  Pricing"]
    row = find_total_annual_cost_row(ws) or 113
    q = quarterly_values(ws, row)
    nz = next((x for x in q if x is not None), 0)
    exl_div = infer_scale_million(nz)
    years, tcv = annualize_quarters(q)
    years, tcv = scale_years_tcv(years, tcv, divide_million=exl_div)
    brief = F4 / "EXL_Evaluation_Brief.docx"
    v = VendorExtract(
        id="exl",
        displayName="EXL",
        color="#EA580C",
        workbookPath=str(path.relative_to(BASE)),
        evaluationBriefText=docx_paragraph_texts(brief),
        flags=[
            "Training (~$46.3M) and dual-run (~$13.1M) are real line items; investment credit (-$19.25M) nets partial offset.",
            "Severance not included (FIS-borne); CCaaS excluded; COLA on actuals for several geographies — TCV not directly comparable to bundled-COLA bids.",
            "Baseline framing in submission uses ~$841M vs FIS ~$750.6M — savings claims not directly comparable.",
        ],
        sources={"totalAnnualCost": f"'6.0  Pricing'!row {row} quarterly C:V (raw $, /1e6)"},
    )
    v.pricing = {
        "totalAnnualCostRow": row,
        "quarterlyUnit": "USD raw → $M" if exl_div else "$M as in workbook (Tab 6.0 C:V)",
        "years": [{"year": i + 1, "valueM": years[i]} for i in range(5)],
        "tcvM": round(tcv, 2),
        "y6M": None,
        "y7M": None,
        "tcv7M": None,
        "provenance": build_pricing_provenance(
            path.name,
            ws,
            row,
            round(tcv, 2),
            years,
            "raw USD → $M (quarterly sums ÷ 1e6)" if exl_div else "values already in $M per workbook cells",
        ),
    }
    v.rateCard = rate_rows_with_snapshots(ws, path.name)
    _patch_exl_cs_live(v.rateCard, ws, path.name)
    v.colaAssumptions = extract_cola_assumptions(ws, "exl")
    eff_ws = wb["9.0 Client EfficiencyAssumption"]
    v.efficiency = extract_efficiency_table(eff_ws, path.name, "9.0 Client EfficiencyAssumption")
    _es = efficiency_snapshot(eff_ws)
    v.efficiency["tableSnapshot"] = _es
    v.efficiency["sourcePreview"] = workbook_meta(
        source_file=path.name,
        tab="9.0 Client EfficiencyAssumption",
        location="Rows 6–12",
        value_label="Efficiency assumptions",
        calculation="Vendor Tab 9.0",
        snapshot=_es,
    )
    v.governance = count_governance(wb["5.0  Technology Solution"])
    v.governanceItems = extract_governance_items(wb["5.0  Technology Solution"])
    for r, lab in [(66, "Setup"), (67, "Rebadging"), (71, "Training"), (72, "Dual run"), (73, "Investment credit")]:
        s = sum_onetime_row(ws, r)
        if s is not None:
            v.oneTimeLines.append(
                {
                    "label": lab,
                    "row": r,
                    "sumQuarterlyUsd": s,
                    "sourcePreview": onetime_source(ws, r, lab, path.name),
                }
            )
    v.migrationNotes = migration_notes_with_waves(wb)
    v.adminTabs = [{"tab": t, "status": "complete" if t in wb.sheetnames else "missing"} for t in STANDARD_TABS]
    v.drilldownSnippets = extract_drill_blocks(wb, drill_ctx)
    enrich_drill_snippets(wb, v.drilldownSnippets, path.name)
    enrich_drill_linked_documents("exl", v.drilldownSnippets)
    return v


def extract_genpact(wb: Optional[openpyxl.Workbook] = None, drill_ctx: Optional[DrillContext] = None) -> VendorExtract:
    path = F1 / "Genpact_Appendix B - TMS RFP - G updated.xlsx"
    if wb is None:
        wb = openpyxl.load_workbook(path, data_only=True)
    if drill_ctx is None:
        drill_ctx = DrillContext(set(), {}, {}, set())
    ws = wb["6.0  Pricing"]
    row = find_total_annual_cost_row(ws) or 131
    q = quarterly_values_extend(ws, row, max_quarters=32)
    years5, tcv5 = annualize_quarters(q[:20])
    years5, tcv5 = scale_years_tcv(years5, tcv5, divide_million=False)
    y6 = y7 = None
    tcv7 = tcv5
    if len(q) >= 24:
        chunk6 = q[20:24]
        if any(x is not None for x in chunk6):
            y6 = round(sum(x for x in chunk6 if x is not None), 2)
    if len(q) >= 28:
        chunk7 = q[24:28]
        if any(x is not None for x in chunk7):
            y7 = round(sum(x for x in chunk7 if x is not None), 2)
    if len(q) > 20:
        tcv7 = round(sum(x for x in q if x is not None), 2)
    brief = F4 / "Genpact_Evaluation_Brief.docx"
    v = VendorExtract(
        id="genpact",
        displayName="Genpact",
        color="#059669",
        workbookPath=str(path.relative_to(BASE)),
        evaluationBriefText=docx_paragraph_texts(brief),
        flags=[
            "7-year minimum term; years 6–7 in continuation quarters on Tab 6.0 row 131 (7-year TCV from workbook sum of all quarters).",
            "Efficiency certainty marked TBD across geographies in Tab 9.0 despite high labor-arb assumption.",
            "Strong governance commit count (verify Tab 5.0); existing FIS rebadge scale called out in materials.",
        ],
        sources={"totalAnnualCost": f"'6.0  Pricing'!row {row} quarters C→AD ($M); 5-year TCV first 20 quarters"},
    )
    v.pricing = {
        "totalAnnualCostRow": row,
        "quarterlyUnit": "$M as submitted",
        "years": [{"year": i + 1, "valueM": round(years5[i], 2) if years5[i] is not None else None} for i in range(5)],
        "tcvM": round(tcv5, 2),
        "y6M": y6,
        "y7M": y7,
        "tcv7M": tcv7,
        "provenance": build_pricing_provenance(
            path.name, ws, row, round(tcv5, 2), years5, "values already in $M per workbook"
        ),
    }
    v.rateCard = rate_rows_with_snapshots(ws, path.name)
    _patch_genpact_cs_live(v.rateCard, ws, path.name)
    v.colaAssumptions = extract_cola_assumptions(ws, "genpact")
    eff_ws = wb["9.0 Client EfficiencyAssumption"]
    v.efficiency = extract_efficiency_table(eff_ws, path.name, "9.0 Client EfficiencyAssumption")
    _eg = efficiency_snapshot(eff_ws)
    v.efficiency["tableSnapshot"] = _eg
    v.efficiency["sourcePreview"] = workbook_meta(
        source_file=path.name,
        tab="9.0 Client EfficiencyAssumption",
        location="Rows 6–12",
        value_label="Efficiency assumptions",
        calculation="Vendor Tab 9.0",
        snapshot=_eg,
    )
    v.governance = count_governance(wb["5.0  Technology Solution"])
    v.governanceItems = extract_governance_items(wb["5.0  Technology Solution"])
    v.migrationNotes = migration_notes_with_waves(wb)
    v.adminTabs = [{"tab": t, "status": "complete" if t in wb.sheetnames else "missing"} for t in STANDARD_TABS]
    v.drilldownSnippets = extract_drill_blocks(wb, drill_ctx)
    enrich_drill_snippets(wb, v.drilldownSnippets, path.name)
    enrich_drill_linked_documents("genpact", v.drilldownSnippets)
    return v


def extract_ibm(wb: Optional[openpyxl.Workbook] = None, drill_ctx: Optional[DrillContext] = None) -> VendorExtract:
    path = F1 / "IBM_Appendix B - TMS RFP - Vendor Response Workbook_IBM 26-Mar-2026.xlsx"
    if wb is None:
        wb = openpyxl.load_workbook(path, data_only=True)
    if drill_ctx is None:
        drill_ctx = DrillContext(set(), {}, {}, set())
    ws = wb["6.0  Pricing"]
    row = find_total_annual_cost_row(ws) or 119
    q = quarterly_values(ws, row)
    years, tcv = annualize_quarters(q)
    years, tcv = scale_years_tcv(years, tcv, divide_million=False)
    # Headline TCV aligned to evaluation brief (~$460.0M); 20Q sum is $459.99M at 2dp
    tcv = round(tcv, 1)
    brief = F4 / "IBM_Evaluation_Brief.docx"
    v = VendorExtract(
        id="ibm",
        displayName="IBM",
        color="#1E40AF",
        workbookPath=str(path.relative_to(BASE)),
        evaluationBriefText=docx_paragraph_texts(brief),
        flags=[
            "Operating TCV ~$460.0M on Row 119; separate IBM investment lines (Rows 68–73, ~$11.3M) are not deductions from Row 119.",
            "COLA shown as 0% pending mutual agreement — treated as pricing floor, not ceiling.",
            "Tab 9.0 efficiency cells largely blank ('vendor to provide'); proposal non-binding language (see brief).",
        ],
        sources={"totalAnnualCost": f"'6.0  Pricing'!row {row} quarterly ($M)"},
    )
    v.pricing = {
        "totalAnnualCostRow": row,
        "quarterlyUnit": "$M as submitted",
        "years": [{"year": i + 1, "valueM": years[i]} for i in range(5)],
        "tcvM": tcv,
        "y6M": None,
        "y7M": None,
        "tcv7M": None,
        "provenance": build_pricing_provenance(
            path.name, ws, row, float(tcv), years, "values in $M per workbook"
        ),
    }
    v.rateCard = rate_rows_with_snapshots(ws, path.name)
    v.colaAssumptions = extract_cola_assumptions(ws, "ibm")
    eff_ws = wb["9.0 Client EfficiencyAssumption"]
    v.efficiency = extract_efficiency_table(eff_ws, path.name, "9.0 Client EfficiencyAssumption")
    _ei = efficiency_snapshot(eff_ws)
    v.efficiency["tableSnapshot"] = _ei
    v.efficiency["sourcePreview"] = workbook_meta(
        source_file=path.name,
        tab="9.0 Client EfficiencyAssumption",
        location="Rows 6–12",
        value_label="Efficiency assumptions",
        calculation="Vendor Tab 9.0",
        snapshot=_ei,
    )
    v.governance = count_governance(wb["5.0  Technology Solution"])
    v.governanceItems = extract_governance_items(wb["5.0  Technology Solution"])
    tr = sum_onetime_row(ws, 67)
    v.oneTimeLines = [
        {
            "label": "FIS Transition (ex row 67)",
            "row": 67,
            "sumQuarterlyUsd": tr,
            "sourcePreview": onetime_source(ws, 67, "FIS Transition", path.name),
        },
        {
            "label": "IBM investment (rows 68–73; separate from operating TCV)",
            "row": 68,
            "sumQuarterlyUsd": None,
            "sourcePreview": onetime_source(ws, 68, "IBM investment (partial row)", path.name),
        },
    ]
    v.migrationNotes = migration_notes_with_waves(wb)
    v.adminTabs = [{"tab": t, "status": "complete" if t in wb.sheetnames else "missing"} for t in STANDARD_TABS]
    v.drilldownSnippets = extract_drill_blocks(wb, drill_ctx)
    enrich_drill_snippets(wb, v.drilldownSnippets, path.name)
    enrich_drill_linked_documents("ibm", v.drilldownSnippets)
    return v


def extract_sutherland(wb: Optional[openpyxl.Workbook] = None, drill_ctx: Optional[DrillContext] = None) -> VendorExtract:
    path = F1 / "Sutherland_Appendix B - TMS RFP - Sutherland Response to FIS.xlsx"
    if wb is None:
        wb = openpyxl.load_workbook(path, data_only=True)
    if drill_ctx is None:
        drill_ctx = DrillContext(set(), {}, {}, set())
    ws = wb["6.0  Pricing"]
    row = find_total_annual_cost_row(ws) or 120
    q = quarterly_values(ws, row)
    nz = next((x for x in q if x is not None), 0)
    sut_div = infer_scale_million(nz)
    years, tcv = annualize_quarters(q)
    years, tcv = scale_years_tcv(years, tcv, divide_million=sut_div)
    eff_name = "9.0 Client EfficiencyAssumption"
    brief = F4 / "Sutherland_Evaluation_Brief.docx"
    v = VendorExtract(
        id="sutherland",
        displayName="Sutherland",
        color="#4B5563",
        workbookPath=str(path.relative_to(BASE)),
        evaluationBriefText=docx_paragraph_texts(brief),
        flags=[
            "TCV in Section 6.7 financial summary (Tab 6.0) — not 'rate card only'.",
            "Certainty % mid-high-30s in Tab 9.0 — lowest explicit certainty among detailed submissions (verify cell).",
        ],
        sources={"totalAnnualCost": f"'6.0  Pricing'!row {row}", "efficiency": f"'{eff_name}' (use tab without '(2)')"},
    )
    v.pricing = {
        "totalAnnualCostRow": row,
        "quarterlyUnit": "USD raw → $M" if sut_div else "$M as in workbook (Tab 6.0 C:V)",
        "years": [{"year": i + 1, "valueM": years[i]} for i in range(5)],
        "tcvM": round(tcv, 2),
        "y6M": None,
        "y7M": None,
        "tcv7M": None,
        "provenance": build_pricing_provenance(
            path.name,
            ws,
            row,
            round(tcv, 2),
            years,
            "raw USD → $M (quarterly sums ÷ 1e6)" if sut_div else "values already in $M per workbook cells",
        ),
    }
    v.rateCard = rate_rows_with_snapshots(ws, path.name)
    _patch_sutherland_cs_live(v.rateCard, ws, path.name)
    v.colaAssumptions = extract_cola_assumptions(ws, "sutherland")
    eff_ws = wb[eff_name]
    v.efficiency = extract_efficiency_table(eff_ws, path.name, eff_name)
    _esu89 = efficiency_snapshot(eff_ws)
    v.efficiency["tableSnapshot"] = _esu89
    v.efficiency["sourcePreview"] = workbook_meta(
        source_file=path.name,
        tab=eff_name,
        location="Rows 6–12",
        value_label="Efficiency assumptions",
        calculation="Vendor Tab 9.0 (tab without duplicate)",
        snapshot=_esu89,
    )
    v.governance = count_governance(wb["5.0  Technology Solution"])
    v.governanceItems = extract_governance_items(wb["5.0  Technology Solution"])
    v.migrationNotes = migration_notes_with_waves(wb)
    v.adminTabs = [{"tab": t, "status": "complete" if t in wb.sheetnames else "missing"} for t in STANDARD_TABS]
    v.drilldownSnippets = extract_drill_blocks(wb, drill_ctx)
    enrich_drill_snippets(wb, v.drilldownSnippets, path.name)
    enrich_drill_linked_documents("sutherland", v.drilldownSnippets)
    return v


def extract_ubiquity(wb: Optional[openpyxl.Workbook] = None, drill_ctx: Optional[DrillContext] = None) -> VendorExtract:
    path = F1 / "Ubiquity_Pricing Response for FIS(2).xlsx"
    if wb is None:
        wb = openpyxl.load_workbook(path, data_only=True)
    if drill_ctx is None:
        drill_ctx = DrillContext(set(), {}, {}, set())
    ws = wb["6.0  Pricing"]
    parsed = parse_ubiquity_annual_from_rows(ws)
    if not parsed:
        parsed = parse_ubiquity_row112(ws)
    years, tcv = parsed if parsed else ([None] * 5, 0.0)
    qpath = F1 / "Ubiquity_Response_Vendor_RFP_Questionnaire_with_Fraud_Supplement_FINAL.xlsx"
    qtext = ""
    ubi_drill: list[dict[str, Any]] = []
    if qpath.exists():
        try:
            qwb = openpyxl.load_workbook(qpath, data_only=True)
            chunks = []
            ubi_drill = extract_questionnaire_blocks(qwb, qpath.name, max_per_sheet=8)
            for sname in qwb.sheetnames[:6]:
                wq = qwb[sname]
                for row in wq.iter_rows(max_row=80, max_col=8, values_only=True):
                    parts = [str(x) for x in row if x and len(str(x).split()) > 8]
                    if parts:
                        chunks.append(" ".join(parts))
            qtext = "\n".join(chunks[:40])
            enrich_drill_snippets(qwb, ubi_drill, qpath.name)
        except Exception:
            pass
    brief = F4 / "Ubiquity_Evaluation_Brief.docx"
    v = VendorExtract(
        id="ubiquity",
        displayName="Ubiquity",
        color="#DC2626",
        workbookPath=str(path.relative_to(BASE)),
        evaluationBriefText=docx_paragraph_texts(brief) + "\n\n[Questionnaire excerpts]\n" + qtext[:8000],
        flags=[
            "Pricing workbook applies rates to current headcount without efficiency ramp — trajectory rises YoY; not a cost-reduction posture vs baseline.",
            "$3M exclusivity payment to FIS and $2M tech migration called out in materials — still well above FIS baseline on TCV.",
        ],
        sources={"totalAnnualCost": "'6.0  Pricing' rows 110+111 quarterly sums (USD → $M) cross-check row 112 annual cells"},
    )
    v.pricing = {
        "totalAnnualCostRow": 112,
        "quarterlyUnit": "USD raw → $M (rows 110+111)",
        "years": [{"year": i + 1, "valueM": years[i]} for i in range(5)],
        "tcvM": round(tcv, 2),
        "y6M": None,
        "y7M": None,
        "tcv7M": None,
        "provenance": ubiquity_price_provenance(path.name, ws, round(tcv, 2), years),
    }
    v.rateCard = rate_rows_with_snapshots(ws, path.name)
    v.colaAssumptions = extract_cola_assumptions(ws, "ubiquity")
    v.migrationNotes = migration_notes_with_waves(wb)
    v.adminTabs = tab_status(wb, STANDARD_TABS)
    v.drilldownSnippets = ubi_drill if ubi_drill else [{"tab": t, "missing": True, "snippets": []} for t in DRILL_TABS]
    enrich_drill_linked_documents("ubiquity", v.drilldownSnippets)
    v.governance = {"commit": 0, "partial": 0, "cannotCommit": 0}
    v.governanceItems = []
    return v


VENDOR_IDS_SC = ["cognizant", "genpact", "exl", "sutherland", "ubiquity", "ibm"]
_NULL_SCORES: dict[str, Any] = {k: None for k in VENDOR_IDS_SC}


def _dim_row(dim_id: str, pillar: str, label: str) -> dict[str, Any]:
    return {"id": dim_id, "pillar": pillar, "label": label, "scores": dict(_NULL_SCORES)}


SCORECARD = {
    "source": "Pending Workshop 1 evaluator sessions — numeric scores intentionally withheld until EOD Thursday, April 3, 2026.",
    "columnOrder": list(VENDOR_IDS_SC),
    "dimensions": [
        _dim_row("1.1", "Commercial Attractiveness (30%)", "1.1 Synergy Alignment"),
        _dim_row("1.2", "Commercial Attractiveness (30%)", "1.2 Efficiency Assumptions & Certainty"),
        _dim_row("1.3", "Commercial Attractiveness (30%)", "1.3 Investments, Transparency & Savings Protection"),
        _dim_row("2.1", "Operational Excellence & Delivery (25%)", "2.1 SLA Commitment Level"),
        _dim_row("2.2", "Operational Excellence & Delivery (25%)", "2.2 Language, Coverage & Compliance"),
        _dim_row("2.3", "Operational Excellence & Delivery (25%)", "2.3 Delivery Model Credibility"),
        _dim_row("3.1", "Technology & AI (15%)", "3.1 Tech Governance Compliance"),
        _dim_row("3.2", "Technology & AI (15%)", "3.2 Migration Plan Specificity"),
        _dim_row("3.3", "Technology & AI (15%)", "3.3 AI Overlay Innovation & Impact"),
        _dim_row("4.1", "Client & Workforce Migration (20%)", "4.1 Wave Composition & Constraint Compliance"),
        _dim_row("4.2", "Client & Workforce Migration (20%)", "4.2 Sequencing, Speed & Revenue Pace"),
        _dim_row("4.3", "Client & Workforce Migration (20%)", "4.3 Workforce Transition Safeguards"),
        _dim_row("5.1", "Partnership Readiness (10%)", "5.1 Reference & Case Study Alignment"),
        _dim_row("5.2", "Partnership Readiness (10%)", "5.2 Operational Depth & Geographic Reach"),
        _dim_row("5.3", "Partnership Readiness (10%)", "5.3 Contracting Posture"),
    ],
    "composite": dict(_NULL_SCORES),
}


def radar_series(vendors: list[dict]) -> dict[str, Any]:
    """Placeholder axes only until Workshop 1 scores populate."""
    zeros = {"commercial": 0.0, "operations": 0.0, "technology": 0.0, "migration": 0.0, "partnership": 0.0}
    ids = SCORECARD["columnOrder"]
    return {"fieldAverage": dict(zeros), "vendors": [{"vendorId": vid, "pillars": dict(zeros)} for vid in ids]}


def normalize_vendor_dict(d: dict[str, Any]) -> None:
    py = d.get("pricing") or {}
    for y in py.get("years") or []:
        if y.get("valueM") is not None:
            y["valueM"] = round(float(y["valueM"]), 2)


def open_all_vendor_workbooks() -> dict[str, Optional[openpyxl.Workbook]]:
    spec = {
        "cognizant": F1 / "Cognizant_ Appendix B - TMS RFP - Cognizant - Vendor Response Workbook.xlsx",
        "exl": F1 / "EXL_APPEND~1.XLS",
        "genpact": F1 / "Genpact_Appendix B - TMS RFP - G updated.xlsx",
        "ibm": F1 / "IBM_Appendix B - TMS RFP - Vendor Response Workbook_IBM 26-Mar-2026.xlsx",
        "sutherland": F1 / "Sutherland_Appendix B - TMS RFP - Sutherland Response to FIS.xlsx",
        "ubiquity": F1 / "Ubiquity_Pricing Response for FIS(2).xlsx",
    }
    out: dict[str, Optional[openpyxl.Workbook]] = {}
    for k, p in spec.items():
        try:
            if not p.exists():
                out[k] = None
                continue
            out[k] = load_workbook_fixed(p) if k == "exl" else openpyxl.load_workbook(str(p), data_only=True)
        except Exception:
            out[k] = None
    return out


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    wbs = open_all_vendor_workbooks()
    qk, cn, cr = (
        build_drill_consensus(wbs)
        if any(wbs.values())
        else (set(), {}, {})
    )
    tmpl = load_template_norm_strings()
    drill_ctx = DrillContext(qk, cn, cr, tmpl)

    rows = [
        extract_cognizant(wbs.get("cognizant"), drill_ctx),
        extract_exl(wbs.get("exl"), drill_ctx),
        extract_genpact(wbs.get("genpact"), drill_ctx),
        extract_ibm(wbs.get("ibm"), drill_ctx),
        extract_sutherland(wbs.get("sutherland"), drill_ctx),
        extract_ubiquity(wbs.get("ubiquity"), drill_ctx),
    ]
    vendors = []
    for v in rows:
        d = asdict(v)
        normalize_vendor_dict(d)
        vendors.append(d)
        (OUT / f"vendor_{v.id}.json").write_text(json.dumps(d, indent=2), encoding="utf-8")

    portfolio = {
        "title": "TMS RFP Intelligence Center",
        "subtitle": "Vendor Memos & Tear Sheets",
        "classification": "FIS x Total Issuing Solutions · Confidential · March 2026",
        "footer": "© 2026 FIS and/or its subsidiaries. All Rights Reserved. FIS confidential and proprietary information.",
        "extractionTimestamp": EXTRACTION_TIMESTAMP,
        "dashboardBuildTimestamp": EXTRACTION_TIMESTAMP,
        "baselineAnnualM": {"low": 144, "high": 156, "mid": 150},
        "synergyTargetM": 40,
        "vendors": [
            {
                "id": x["id"],
                "displayName": x["displayName"],
                "color": x["color"],
                "tcvM": x["pricing"]["tcvM"],
                "composite": None,
            }
            for x in vendors
        ],
        "scorecard": SCORECARD,
        "radar": radar_series(vendors),
    }
    (OUT / "portfolio.json").write_text(json.dumps(portfolio, indent=2), encoding="utf-8")
    (OUT / "scorecard.json").write_text(json.dumps(SCORECARD, indent=2), encoding="utf-8")
    print("Wrote", len(vendors), "vendors + portfolio to", OUT)
    _sync_data_json_to_public_and_out()


def _sync_data_json_to_public_and_out() -> None:
    """Copy src/data JSON to public/data (Next dev) and out/data (static export) so the UI can load updates without rebuilding."""
    pub = BASE / "public" / "data"
    pub.mkdir(parents=True, exist_ok=True)
    for path in [OUT / "portfolio.json", OUT / "scorecard.json", *OUT.glob("vendor_*.json")]:
        if path.is_file():
            shutil.copy2(path, pub / path.name)
    # Dashboard JSON produced by other scripts (workshop memos, scores import, ideal RFP build).
    for name in ("workshop1_memos.json", "idealRfpSubmission.json", "evaluatorScores.json"):
        p = OUT / name
        if p.is_file():
            shutil.copy2(p, pub / name)
    out_data = BASE / "out" / "data"
    if (BASE / "out").is_dir():
        out_data.mkdir(parents=True, exist_ok=True)
        for f in pub.glob("*.json"):
            shutil.copy2(f, out_data / f.name)
        print("Synced data JSON to", pub.relative_to(BASE), "and", out_data.relative_to(BASE))


if __name__ == "__main__":
    main()
